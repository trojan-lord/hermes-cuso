#!/usr/bin/env python3
"""
Professional Manuscript Formatter
==================================
Converts plain-text chapter input into a publication-ready .docx manuscript.

Industry-standard formatting:
  - Times New Roman 12pt throughout
  - Double line spacing
  - 1-inch margins (top, bottom, left, right)
  - Paragraphs indented 0.5 inches, no extra inter-paragraph spacing
  - Chapter headings: centered, bold, each on a new page
  - Scene breaks: three centered asterisks (* * *)
  - Title page: title ~1/3 down, author below, word count at bottom

Usage:
    python format_manuscript.py input.txt -o output.docx \
        --title "My Novel" --author "Jane Doe"
"""

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
MARGIN = Inches(1)
PARAGRAPH_INDENT = Inches(0.5)
CHAPTER_HEADING_SIZE = Pt(16)
SCENE_BREAK_TEXT = "* * *"


@dataclass
class Chapter:
    title: str
    paragraphs: List[str] = field(default_factory=list)


@dataclass
class ManuscriptMeta:
    title: str = "Untitled"
    author: str = "Author Name"
    subtitle: str = ""
    contact: str = ""


def _is_chapter_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if re.match(r"^chapter\s+\d+\s*:", stripped, re.IGNORECASE):
        return True
    if stripped.startswith("# ") and re.match(r"#\s+chapter\s+\d+\s*:", stripped, re.IGNORECASE):
        return True
    return False


def _is_scene_break(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if re.fullmatch(r"[\s*]+", stripped) and "*" in stripped:
        return True
    return False


def _clean_heading(line: str) -> str:
    return line.strip().lstrip("#").strip()


def parse_text(text: str) -> List[Chapter]:
    lines = text.splitlines()
    chapters: List[Chapter] = []
    current: Optional[Chapter] = None
    buffer: List[str] = []

    def _flush_paragraph():
        nonlocal buffer
        if buffer and current is not None:
            para = " ".join(buffer).strip()
            if para:
                current.paragraphs.append(para)
            buffer = []

    for line in lines:
        stripped = line.strip()

        if _is_chapter_heading(stripped):
            _flush_paragraph()
            if current is not None:
                chapters.append(current)
            current = Chapter(title=_clean_heading(stripped))
            continue

        if _is_scene_break(stripped):
            _flush_paragraph()
            if current is not None:
                current.paragraphs.append(SCENE_BREAK_TEXT)
            continue

        if not stripped:
            _flush_paragraph()
            continue

        if current is None:
            current = Chapter(title="")
        buffer.append(stripped)

    _flush_paragraph()
    if current is not None:
        chapters.append(current)

    if not chapters:
        chapters = [Chapter(title="", paragraphs=[text.strip()])]

    return chapters


def count_words(chapters: List[Chapter]) -> int:
    total = 0
    for ch in chapters:
        for para in ch.paragraphs:
            if para == SCENE_BREAK_TEXT:
                continue
            total += len(para.split())
    return total


def _set_font(run):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


def _set_double_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _configure_section(section):
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)


def _new_page(document):
    section = document.add_section()
    _configure_section(section)
    return section


def _build_title_page(doc, meta):
    for _ in range(12):
        p = doc.add_paragraph()
        _set_double_spacing(p)
        _set_font(p.add_run(""))

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_double_spacing(title_para)
    run = title_para.add_run(meta.title)
    _set_font(run)
    run.bold = True
    run.font.size = Pt(24)

    if meta.subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_double_spacing(sub_para)
        run = sub_para.add_run(meta.subtitle)
        _set_font(run)
        run.font.size = Pt(16)

    spacer = doc.add_paragraph()
    _set_double_spacing(spacer)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_double_spacing(author_para)
    run = author_para.add_run(meta.author)
    _set_font(run)
    run.font.size = Pt(14)

    if meta.contact:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_double_spacing(contact_para)
        run = contact_para.add_run(meta.contact)
        _set_font(run)
        run.font.size = Pt(10)

    for _ in range(12):
        p = doc.add_paragraph()
        _set_double_spacing(p)
        _set_font(p.add_run(""))


def _render_chapter(doc, chapter, is_first=False):
    if not is_first:
        _new_page(doc)

    if chapter.title:
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_double_spacing(heading_para)
        heading_para.paragraph_format.space_before = Pt(0)
        run = heading_para.add_run(chapter.title)
        _set_font(run)
        run.bold = True
        run.font.size = CHAPTER_HEADING_SIZE

        spacer = doc.add_paragraph()
        _set_double_spacing(spacer)

    for para_text in chapter.paragraphs:
        if para_text == SCENE_BREAK_TEXT:
            sb = doc.add_paragraph()
            sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_double_spacing(sb)
            run = sb.add_run(SCENE_BREAK_TEXT)
            _set_font(run)
        else:
            p = doc.add_paragraph()
            _set_double_spacing(p)
            p.paragraph_format.first_line_indent = PARAGRAPH_INDENT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(para_text)
            _set_font(run)


def format_manuscript(input_path, output_path, title="Untitled", author="Author Name", subtitle="", contact=""):
    text = Path(input_path).read_text(encoding="utf-8")
    chapters = parse_text(text)
    meta = ManuscriptMeta(title=title, author=author, subtitle=subtitle, contact=contact)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    _configure_section(doc.sections[0])
    _build_title_page(doc, meta)

    for idx, chapter in enumerate(chapters):
        _render_chapter(doc, chapter, is_first=(idx == 0))

    wc = count_words(chapters)
    wc_para = doc.add_paragraph()
    wc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_double_spacing(wc_para)
    run = wc_para.add_run(f"(Approximately {wc:,} words)")
    _set_font(run)
    run.font.size = Pt(10)
    run.italic = True

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def main():
    parser = argparse.ArgumentParser(description="Format a plain-text manuscript into a professional .docx file.")
    parser.add_argument("input", help="Path to the plain-text manuscript file")
    parser.add_argument("-o", "--output", default=None, help="Output .docx path")
    parser.add_argument("--title", default=None, help="Book title")
    parser.add_argument("--author", default=None, help="Author name")
    parser.add_argument("--subtitle", default="", help="Subtitle (optional)")
    parser.add_argument("--contact", default="", help="Contact info (optional)")
    parser.add_argument("--meta", default=None, help="Path to JSON metadata file")
    args = parser.parse_args()

    title = args.title or "Untitled"
    author = args.author or "Author Name"

    if args.meta:
        with open(args.meta) as f:
            meta = json.load(f)
        title = meta.get("title", title)
        author = meta.get("author", author)

    output = args.output or str(Path(args.input).with_suffix(".docx"))
    result = format_manuscript(args.input, output, title=title, author=author, subtitle=args.subtitle, contact=args.contact)
    print(f"Manuscript formatted: {result}")


if __name__ == "__main__":
    main()
