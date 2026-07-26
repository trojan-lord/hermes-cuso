#!/usr/bin/env python3
"""
Professional Manuscript Formatter (v2)
=======================================
Industry-standard formatting per book-editing skill.

Standards applied:
  • Courier 12pt (submission standard)
  • Double line spacing, NO extra space between paragraphs
  • 1-inch margins, 8.5" x 11" US Letter
  • 0.5-inch first-line paragraph indent
  • Left-justified (not full-justified)
  • Title page: title ~1/3 down, author, word count at bottom
  • Chapter headings: centered, bold, each on new page
  • Scene breaks: centered # or * * *
  • Page numbers at bottom center
  • "THE END" marker after final chapter

Usage:
    python format_manuscript_v2.py input.txt -o output.docx \
        --title "My Novel" --author "Jane Doe"
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips


# ---------------------------------------------------------------------------
# Constants — per book-editing skill
# ---------------------------------------------------------------------------
FONT_NAME = "Courier New"          # Submission standard
FONT_SIZE = Pt(12)
MARGIN = Inches(1)
PARAGRAPH_INDENT = Inches(0.5)
CHAPTER_HEADING_SIZE = Pt(14)      # Slightly larger than body, still Courier
SCENE_BREAK_TEXT = "* * *"
PAGE_NUMBER_SIZE = Pt(10)


# ---------------------------------------------------------------------------
# Data helpers
# ---------------------------------------------------------------------------
@dataclass
class Chapter:
    title: str
    paragraphs: List[str] = field(default_factory=list)


@dataclass
class ManuscriptMeta:
    title: str = "Untitled"
    author: str = "Author Name"
    subtitle: str = ""
    contact: str = ""
    word_count: int = 0


# ---------------------------------------------------------------------------
# Parsing — stricter chapter detection
# ---------------------------------------------------------------------------
def _is_chapter_heading(line: str) -> bool:
    """Only 'Chapter N: Title' or '# Chapter N: Title'."""
    stripped = line.strip()
    if not stripped:
        return False
    if re.match(r"^chapter\s+\d+\s*:", stripped, re.IGNORECASE):
        return True
    if stripped.startswith("# ") and re.match(r"#\s+chapter\s+\d+\s*:", stripped, re.IGNORECASE):
        return True
    return False


def _is_scene_break(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if re.fullmatch(r"[\s*#]+", stripped) and ("*" in stripped or "#" in stripped):
        if len(stripped) <= 10:
            return True
    return False


def _clean_heading(line: str) -> str:
    return line.strip().lstrip("#").strip()


def parse_text(text: str) -> List[Chapter]:
    lines = text.splitlines()
    chapters: List[Chapter] = []
    current: Optional[Chapter] = None
    buffer: List[str] = []

    def _flush_paragraph():
        nonlocal buffer
        if buffer and current is not None:
            para = " ".join(buffer).strip()
            if para:
                current.paragraphs.append(para)
            buffer = []

    for line in lines:
        stripped = line.strip()

        if _is_chapter_heading(stripped):
            _flush_paragraph()
            if current is not None:
                chapters.append(current)
            current = Chapter(title=_clean_heading(stripped))
            continue

        if _is_scene_break(stripped):
            _flush_paragraph()
            if current is not None:
                current.paragraphs.append(SCENE_BREAK_TEXT)
            continue

        if not stripped:
            _flush_paragraph()
            continue

        if current is None:
            current = Chapter(title="")
        buffer.append(stripped)

    _flush_paragraph()
    if current is not None:
        chapters.append(current)

    if not chapters:
        chapters = [Chapter(title="", paragraphs=[text.strip()])]

    return chapters


# ---------------------------------------------------------------------------
# Word count
# ---------------------------------------------------------------------------
def count_words(chapters: List[Chapter]) -> int:
    total = 0
    for ch in chapters:
        for para in ch.paragraphs:
            if para == SCENE_BREAK_TEXT:
                continue
            total += len(para.split())
    return total


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------
def _set_font(run, font_name=FONT_NAME, font_size=FONT_SIZE, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _set_double_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _configure_section(section):
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)


def _new_page(document: Document):
    section = document.add_section()
    _configure_section(section)
    return section


def _add_page_numbers(doc: Document):
    """Add page numbers at bottom center of each section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_double_spacing(p)
        
        # Add PAGE field
        run = p.add_run()
        _set_font(run, font_size=PAGE_NUMBER_SIZE)
        fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fldChar1)
        
        run2 = p.add_run()
        _set_font(run2, font_size=PAGE_NUMBER_SIZE)
        instrText = run2._element.makeelement(qn("w:instrText"), {})
        instrText.text = " PAGE "
        run2._element.append(instrText)
        
        run3 = p.add_run()
        _set_font(run3, font_size=PAGE_NUMBER_SIZE)
        fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._element.append(fldChar2)


# ---------------------------------------------------------------------------
# Title page — per industry standard
# ---------------------------------------------------------------------------
def _build_title_page(doc: Document, meta: ManuscriptMeta):
    # Push title to ~1/3 down page
    for _ in range(10):
        p = doc.add_paragraph()
        _set_double_spacing(p)
        _set_font(p.add_run(""))

    # Contact info (top left if provided)
    if meta.contact:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_double_spacing(contact_para)
        run = contact_para.add_run(meta.contact)
        _set_font(run, font_size=Pt(10))

    # Title — centered, bold
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_double_spacing(title_para)
    run = title_para.add_run(meta.title)
    _set_font(run, bold=True, font_size=Pt(18))

    # Subtitle (optional)
    if meta.subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_double_spacing(sub_para)
        run = sub_para.add_run(meta.subtitle)
        _set_font(run, font_size=Pt(14))

    # Spacer
    spacer = doc.add_paragraph()
    _set_double_spacing(spacer)

    # Author — centered
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_double_spacing(author_para)
    run = author_para.add_run(meta.author)
    _set_font(run, font_size=Pt(14))

    # Push word count to bottom
    for _ in range(18):
        p = doc.add_paragraph()
        _set_double_spacing(p)
        _set_font(p.add_run(""))


# ---------------------------------------------------------------------------
# Chapter rendering
# ---------------------------------------------------------------------------
def _render_chapter(doc: Document, chapter: Chapter, is_first: bool = False):
    # New page for each chapter
    if not is_first:
        _new_page(doc)

    # Chapter heading — centered, bold
    if chapter.title:
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_double_spacing(heading_para)
        heading_para.paragraph_format.space_before = Pt(0)
        run = heading_para.add_run(chapter.title)
        _set_font(run, bold=True, font_size=CHAPTER_HEADING_SIZE)

        # 3 blank lines before first paragraph (per skill standard)
        for _ in range(3):
            spacer = doc.add_paragraph()
            _set_double_spacing(spacer)

    # Paragraphs
    for para_text in chapter.paragraphs:
        if para_text == SCENE_BREAK_TEXT:
            sb = doc.add_paragraph()
            sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_double_spacing(sb)
            run = sb.add_run(SCENE_BREAK_TEXT)
            _set_font(run)
        else:
            p = doc.add_paragraph()
            _set_double_spacing(p)
            p.paragraph_format.first_line_indent = PARAGRAPH_INDENT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-justified per standard
            run = p.add_run(para_text)
            _set_font(run)


# ---------------------------------------------------------------------------
# THE END marker
# ---------------------------------------------------------------------------
def _render_end(doc: Document):
    _new_page(doc)
    for _ in range(12):
        p = doc.add_paragraph()
        _set_double_spacing(p)
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_double_spacing(end_para)
    run = end_para.add_run("THE END")
    _set_font(run, bold=True, font_size=Pt(14))


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------
def format_manuscript(
    input_path: str | Path,
    output_path: str | Path,
    title: str = "Untitled",
    author: str = "Author Name",
    subtitle: str = "",
    contact: str = "",
) -> Path:
    text = Path(input_path).read_text(encoding="utf-8")
    chapters = parse_text(text)
    wc = count_words(chapters)
    meta = ManuscriptMeta(
        title=title, author=author, subtitle=subtitle,
        contact=contact, word_count=wc
    )

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Configure first section
    _configure_section(doc.sections[0])

    # Title page
    _build_title_page(doc, meta)

    # Chapters
    for idx, chapter in enumerate(chapters):
        _render_chapter(doc, chapter, is_first=(idx == 0))

    # THE END
    _render_end(doc)

    # Word count on title page (appended at end of first section)
    wc_para = doc.add_paragraph()
    wc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_double_spacing(wc_para)
    run = wc_para.add_run(f"(Approximately {wc:,} words)")
    _set_font(run, font_size=Pt(10), italic=True)

    # Page numbers
    _add_page_numbers(doc)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Format a plain-text manuscript into a professional .docx file.",
    )
    parser.add_argument("input", help="Path to the plain-text manuscript file")
    parser.add_argument("-o", "--output", default=None, help="Output .docx path")
    parser.add_argument("--title", default=None, help="Book title")
    parser.add_argument("--author", default=None, help="Author name")
    parser.add_argument("--subtitle", default="", help="Subtitle (optional)")
    parser.add_argument("--contact", default="", help="Contact info for title page")
    parser.add_argument("--meta", default=None, help="Path to JSON metadata file")

    args = parser.parse_args()

    title = args.title or "Untitled"
    author = args.author or "Author Name"
    subtitle = args.subtitle
    contact = args.contact

    if args.meta:
        with open(args.meta) as f:
            meta = json.load(f)
        title = meta.get("title", title)
        author = meta.get("author", author)
        subtitle = meta.get("subtitle", subtitle)
        contact = meta.get("contact", contact)

    output = args.output or str(Path(args.input).with_suffix(".docx"))

    result = format_manuscript(
        input_path=args.input,
        output_path=output,
        title=title,
        author=author,
        subtitle=subtitle,
        contact=contact,
    )
    print(f"✓ Manuscript formatted: {result}")


if __name__ == "__main__":
    main()
